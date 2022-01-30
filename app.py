import os
from flask import Flask, request, redirect, url_for, render_template, send_from_directory
from werkzeug.utils import secure_filename
# from PyPDF2 import PdfFileReader, PdfFileWriter
import pandas as pd 
import pickle
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.naive_bayes import MultinomialNB
import matplotlib.image as mpimg
import PIL.Image
import matplotlib
from io import StringIO
import matplotlib.pyplot as plt
from sklearn.feature_extraction.text import TfidfVectorizer
from wordcloud import WordCloud
from nltk.corpus import stopwords
from nltk import word_tokenize
import warnings
warnings.filterwarnings('ignore')


clf=pickle.load(open("nlp_model.pkl",'rb'))
cv=pickle.load(open('tranform.pkl','rb'))

UPLOAD_FOLDER = os.path.dirname(os.path.abspath(__file__)) + '/uploads/'
DOWNLOAD_FOLDER = os.path.dirname(os.path.abspath(__file__)) + '/downloads/'
ALLOWED_EXTENSIONS = {'csv','txt'}

app = Flask(__name__, static_url_path="/static")
DIR_PATH = os.path.dirname(os.path.realpath(__file__))
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['DOWNLOAD_FOLDER'] = DOWNLOAD_FOLDER
# limit upload size upto 8mb
app.config['MAX_CONTENT_LENGTH'] = 8 * 1024 * 1024


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        if 'file' not in request.files:
            print('No file attached in request')
            return redirect(request.url)
        file = request.files['file']
        if file.filename == '':
            print('No file selected')
            return redirect(request.url)
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
            process_file(os.path.join(app.config['UPLOAD_FOLDER'], filename), filename)
            return redirect(url_for('uploaded_file', filename=filename))
    return render_template('index.html')

@app.route('/uploads/<filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['DOWNLOAD_FOLDER'],filename+'.docx', as_attachment=True)
    # return send_from_directory(app.config['DOWNLOAD_FOLDER'], filename, as_attachment=True)


def process_file(path, filename):
    file=filename
    # remove_watermark(path, filename)
    test = pd.read_csv(path)
    testing_review=test["rating"].tolist()
    test_data_features = cv.transform(testing_review)
    predictions_test = clf.predict(test_data_features)
    output = pd.DataFrame(data={"ID":test["id"],"rating":test["rating"], "prediction":predictions_test} )
    output.to_csv( app.config['DOWNLOAD_FOLDER']+filename,index=False)
    doc2pdf(filename)

def doc2pdf(file):

    df = pd.read_csv( app.config['DOWNLOAD_FOLDER']+file)
    df.prediction.value_counts().plot(kind='pie', autopct='%1.0f%%', colors=["azure", "lightgrey", "thistle","khaki","pink"])
    plt.savefig("totalpercent")
    # Barplot of all teacher
    faculty_sentiment = df.groupby(['ID', 'prediction']).prediction.count().unstack()
    faculty_sentiment.plot(kind='bar')
    plt.savefig("com_faculty")

    convert_dict = {'ID':str } 
    df = df.astype(convert_dict)
    faculty_ID=list(df['ID'].unique()) 


    for i in faculty_ID:
        plt.figure()
        var = df.loc[(df['ID']== f"{i}")]
        var.prediction.value_counts().plot(kind='pie',autopct='%1.0f%%')
        plt.title(f'{i} faculty')
        plt.savefig(f"{i}",dpi=300)
        # plt.show()


    from docx import Document
    from docx.shared import Inches
    from docx.shared import Pt
    document = Document()

    document.add_heading('Faculty Review Analysis ', 0)

    document.add_heading('Distribution of predicted sentiments ', level=1)
    document.add_picture('totalpercent.png', width=Inches(4))

    lis = []
    for i in faculty_ID:
        var = df.loc[(df['ID']== f"{i}")]
        print("name of faulty "+ i)
        a=(var['prediction'].value_counts())
        a=list(a)
        document.add_heading(f'Faculty ID {i}', level=1)
        document.add_paragraph(f'Number of positive feedback :{a[0]}' , style='List Bullet')
        document.add_paragraph(f'Number of negative feedback :{a[1]}' , style='List Bullet')
        document.add_paragraph(f'Number of netural feedback :{a[2]}' , style='List Bullet')
        document.add_paragraph('')
        document.add_picture(f'{i}.png', width=Inches(4))
        document.add_paragraph('')

    word_cloud_text = ''.join(df['rating'])

    wordcloud = WordCloud(max_font_size=100, # Maximum font size for the largest word
                      max_words=100, # The maximum number of words
                      background_color="white", # Background color for the word cloud image
                      scale = 10, # Scaling between computation and drawing
                      width=800, # Width of the canvas
                      height=400 # Height of the canvas
                     ).generate(word_cloud_text)

    plt.figure()
    plt.imshow(wordcloud, 
            interpolation="bilinear") # to make the displayed image appear more smoothly
    plt.axis("off")
    # plt.show()
    wordcloud.to_file('features_sentiments.png')


    document.add_heading('Students Sentiments in wordcloud ', level=1)
    document.add_picture('features_sentiments.png', width=Inches(5))
    document.add_paragraph('')
    document.add_heading('Comparasion of among all Faculty ', level=1)
    document.add_picture('com_faculty.png', width=Inches(5))
    document.save(app.config['DOWNLOAD_FOLDER']+file+'.docx')

@app.route('/predict',methods=['POST'])
def predict():

	pred="  "
	if request.method == 'POST':
		message = request.form['message']
		data = [message]
		vect = cv.transform(data).toarray()
		my_prediction = clf.predict(vect)
		print(my_prediction)
		if my_prediction =="positive":
			pred='Positive'
		if my_prediction=="negative":
			pred='Negative'
		if my_prediction=="netural":
			pred='Netural'
		 	
	return render_template('result.html',prediction = pred)




if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    app.run(host='0.0.0.0', port=port)
